from io import BytesIO
from pathlib import Path
from typing import Iterable, List, Tuple
import docx
import pdfplumber
import PyPDF2

def _read_pdf(path: Path) -> str:
    out = []
    try:
        with pdfplumber.open(path) as pdf:
            for pg in pdf.pages:
                t = pg.extract_text()
                if t:
                    out.append(t)
        if out:
            return "\n".join(out)
    except:
        pass
    try:
        with open(path, "rb") as f:
            rdr = PyPDF2.PdfReader(f)
            for pg in rdr.pages:
                t = pg.extract_text()
                if t:
                    out.append(t)
    except:
        pass
    return "\n".join(out) if out else ""

def _read_docx(path: Path) -> str:
    try:
        doc = docx.Document(path)
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for tbl in doc.tables:
            for row in tbl.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    except:
        return ""

def _from_bytes(name: str, data: bytes) -> str:
    ext = Path(name).suffix.lower()
    if ext == ".pdf":
        try:
            with pdfplumber.open(BytesIO(data)) as pdf:
                txt = "\n".join(pg.extract_text() for pg in pdf.pages if pg.extract_text())
            if txt.strip():
                return txt
        except:
            pass
        try:
            rdr = PyPDF2.PdfReader(BytesIO(data))
            txt = "\n".join(pg.extract_text() for pg in rdr.pages if pg.extract_text())
            return txt
        except:
            return ""
    elif ext == ".docx":
        try:
            doc = docx.Document(BytesIO(data))
            parts = [p.text for p in doc.paragraphs if p.text.strip()]
            for tbl in doc.tables:
                for row in tbl.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))
            return "\n".join(parts)
        except:
            return ""
    elif ext == ".txt":
        return data.decode("utf-8", errors="ignore")
    return ""

def _from_file(path: Path) -> str:
    ext = path.suffix.lower()
    if ext == ".pdf":
        return _read_pdf(path)
    elif ext == ".docx":
        return _read_docx(path)
    elif ext == ".txt":
        return path.read_text(encoding="utf-8")
    return ""

def load_resumes(dir_path: str) -> Tuple[List[str], List[str]]:
    p = Path(dir_path)
    files = sorted(f for f in p.iterdir() if f.suffix.lower() in {".txt", ".pdf", ".docx"})
    ids = [f.stem for f in files]
    texts = [_from_file(f).strip() for f in files]
    return ids, texts

def load_uploaded(files: Iterable[Tuple[str, bytes]]) -> Tuple[List[str], List[str]]:
    ids, texts = [], []
    for fname, data in files:
        cid = Path(fname).stem
        txt = _from_bytes(fname, data).strip()
        ids.append(cid)
        texts.append(txt)
    return ids, texts