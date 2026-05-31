from io import BytesIO
from pathlib import Path
from typing import Iterable, List, Tuple
import docx
import pdfplumber
import PyPDF2


def read_pdf_content(file_path: Path) -> str:
    content = ""
    
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    content += page_text + "\n"
        if content.strip():
            return content
    except Exception as e:
        print(f"[pdfplumber] Error: {e}")
    
    try:
        with open(file_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    content += page_text + "\n"
    except Exception as e:
        print(f"[PyPDF2] Error: {e}")
    
    return content


def read_docx_content(file_path: Path) -> str:
    try:
        doc = docx.Document(file_path)
        parts = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        
        return "\n".join(parts)
    except Exception as e:
        print(f"[DOCX] Error: {e}")
        return ""


def extract_from_bytes(filename: str, raw_data: bytes) -> str:
    ext = Path(filename).suffix.lower()
    
    if ext == ".pdf":
        content = ""
        try:
            with pdfplumber.open(BytesIO(raw_data)) as pdf:
                for page in pdf.pages:
                    txt = page.extract_text()
                    if txt:
                        content += txt + "\n"
            if content.strip():
                return content
        except Exception as e:
            print(f"[pdfplumber] Error: {e}")
        
        try:
            reader = PyPDF2.PdfReader(BytesIO(raw_data))
            for page in reader.pages:
                txt = page.extract_text()
                if txt:
                    content += txt + "\n"
        except Exception as e:
            print(f"[PyPDF2] Error: {e}")
        return content
    
    elif ext == ".docx":
        try:
            doc = docx.Document(BytesIO(raw_data))
            parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    parts.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))
            return "\n".join(parts)
        except Exception as e:
            print(f"[DOCX] Error: {e}")
        return ""
    
    elif ext == ".txt":
        return raw_data.decode("utf-8", errors="ignore")
    
    return ""


def extract_from_file(file_path: Path) -> str:
    ext = file_path.suffix.lower()
    if ext == ".pdf":
        return read_pdf_content(file_path)
    elif ext == ".docx":
        return read_docx_content(file_path)
    elif ext == ".txt":
        return file_path.read_text(encoding="utf-8")
    return ""


def load_resumes(directory: str) -> Tuple[List[str], List[str]]:
    path = Path(directory)
    files = sorted(p for p in path.iterdir() if p.suffix.lower() in {".txt", ".pdf", ".docx"})
    ids = [p.stem for p in files]
    texts = [extract_from_file(p).strip() for p in files]
    return ids, texts


def load_uploaded(files: Iterable[Tuple[str, bytes]]) -> Tuple[List[str], List[str]]:
    ids, texts = [], []
    for fname, data in files:
        cid = Path(fname).stem
        txt = extract_from_bytes(fname, data).strip()
        ids.append(cid)
        texts.append(txt)
    return ids, texts