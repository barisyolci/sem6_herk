from io import BytesIO
from pathlib import Path
from typing import Iterable, List, Tuple

import docx
import pdfplumber
import PyPDF2


def extract_text_from_pdf(file_path: Path) -> str:
    text = ""

    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        if text.strip():
            return text
    except Exception as exc:
        print(f"pdfplumber failed: {exc}")

    try:
        with open(file_path, "rb") as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as exc:
        print(f"PyPDF2 failed: {exc}")

    return text


def extract_text_from_docx(file_path: Path) -> str:
    try:
        doc = docx.Document(file_path)
        text_parts = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    text_parts.append(" | ".join(row_text))

        return "\n".join(text_parts)
    except Exception as exc:
        print(f"DOCX extraction failed: {exc}")
        return ""


def extract_text_from_bytes(filename: str, data: bytes) -> str:
    file_ext = Path(filename).suffix.lower()

    if file_ext == ".pdf":
        text = ""
        try:
            with pdfplumber.open(BytesIO(data)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            if text.strip():
                return text
        except Exception as exc:
            print(f"pdfplumber failed: {exc}")

        try:
            pdf_reader = PyPDF2.PdfReader(BytesIO(data))
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        except Exception as exc:
            print(f"PyPDF2 failed: {exc}")
        return text

    if file_ext == ".docx":
        try:
            doc = docx.Document(BytesIO(data))
            text_parts = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        text_parts.append(" | ".join(row_text))

            return "\n".join(text_parts)
        except Exception as exc:
            print(f"DOCX extraction failed: {exc}")
            return ""

    if file_ext == ".txt":
        return data.decode("utf-8", errors="ignore")

    return ""


def extract_text_from_file(file_path: Path) -> str:
    file_ext = file_path.suffix.lower()

    if file_ext == ".pdf":
        return extract_text_from_pdf(file_path)
    if file_ext == ".docx":
        return extract_text_from_docx(file_path)
    if file_ext == ".txt":
        return file_path.read_text(encoding="utf-8")
    return ""


def load_cvs(cvs_dir: str) -> Tuple[List[str], List[str]]:
    directory = Path(cvs_dir)
    allowed = {".txt", ".pdf", ".docx"}
    files = sorted(path for path in directory.iterdir() if path.suffix.lower() in allowed)
    cv_ids = [path.stem for path in files]
    texts = [extract_text_from_file(path).strip() for path in files]
    return cv_ids, texts


def load_uploaded_files(files: Iterable[Tuple[str, bytes]]) -> Tuple[List[str], List[str]]:
    cv_ids: List[str] = []
    texts: List[str] = []
    for filename, data in files:
        cv_id = Path(filename).stem
        text = extract_text_from_bytes(filename, data).strip()
        cv_ids.append(cv_id)
        texts.append(text)
    return cv_ids, texts
